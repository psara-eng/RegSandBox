from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import PyPDF2
import docx
from bs4 import BeautifulSoup
import io
from emergentintegrations.llm.chat import LlmChat, UserMessage
import json
import base64
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
import csv

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# LLM Configuration
EMERGENT_LLM_KEY = os.environ.get('EMERGENT_LLM_KEY')

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_type: str
    file_size: int
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    parse_status: str = "pending"  # pending, processing, completed, failed
    total_statements: int = 0
    file_content_base64: Optional[str] = None

class Statement(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    sys_id: str = Field(default_factory=lambda: str(uuid.uuid4()))  # Immutable internal ID
    document_id: str
    hierarchy_path: str
    section_ref: str
    section_title: Optional[str] = None
    page_number: Optional[int] = None
    regulation_text: str
    statement_type: str  # Obligation, Prohibition, Recommendation, Definition, Exception
    custom_fields: Dict[str, Any] = Field(default_factory=dict)
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    
    # Restructuring fields
    parent_sys_id: Optional[str] = None
    user_edit_kind: str = "original"  # original, split_child, merge_result, group_parent
    provenance: Dict[str, Any] = Field(default_factory=lambda: {
        "source_sys_ids": [],
        "source_span": None,
        "op_history": []
    })
    user_section_ref: Optional[str] = None
    lock_original_fields: bool = True
    order_index: float = 0.0
    is_superseded: bool = False

class ColumnDefinition(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    name: str
    column_type: str  # text, long_text, number, date, enum, multi_select, checkbox, url, formula
    options: Optional[List[str]] = None
    default_value: Optional[Any] = None
    validation_rules: Optional[Dict[str, Any]] = None
    is_visible: bool = True
    order: int = 0

class Template(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    name: str
    description: Optional[str] = None
    columns: List[Dict[str, Any]]
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class StatementUpdate(BaseModel):
    custom_fields: Dict[str, Any]

class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    message: str

# Restructuring Models
class SplitRequest(BaseModel):
    base_sys_id: str
    splits: List[Dict[str, Any]]  # [{start, end, user_section_ref?}, ...]
    inherit_user_cols: bool = False

class MergeRequest(BaseModel):
    sys_ids: List[str]
    delimiter: str = "\n"
    user_section_ref: Optional[str] = None

class GroupRequest(BaseModel):
    title: str
    sys_ids: List[str]

class ReorderRequest(BaseModel):
    parent_sys_id: Optional[str]
    ordered_sys_ids: List[str]

class UndoRedoState(BaseModel):
    operation: str
    timestamp: datetime
    affected_sys_ids: List[str]
    snapshot_data: Dict[str, Any]

# Helper Functions
async def extract_text_from_pdf(file_content: bytes) -> tuple[str, int]:
    """Extract text from PDF and return text with page count"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text += f"\n[PAGE {page_num}]\n"
            text += page.extract_text()
        return text, len(pdf_reader.pages)
    except Exception as e:
        logging.error(f"PDF extraction error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract PDF text: {str(e)}")

async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text, 0
    except Exception as e:
        logging.error(f"DOCX extraction error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract DOCX text: {str(e)}")

async def extract_text_from_html(file_content: bytes) -> str:
    """Extract text from HTML"""
    try:
        soup = BeautifulSoup(file_content, 'html.parser')
        text = soup.get_text(separator='\n')
        return text, 0
    except Exception as e:
        logging.error(f"HTML extraction error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract HTML text: {str(e)}")

async def parse_with_ai(raw_text: str, document_id: str) -> List[Dict[str, Any]]:
    """Use AI to extract structured statements from raw text"""
    try:
        chat = LlmChat(
            api_key=EMERGENT_LLM_KEY,
            session_id=f"doc_{document_id}",
            system_message="""You are an expert regulatory document parser. Extract atomic statements from regulatory text.
            For each statement, identify:
            1. hierarchy_path: The full path (e.g., "Chapter 3 > Section 10 > 10.2(b)")
            2. section_ref: The specific reference (e.g., "10.2(b)")
            3. section_title: The title if present
            4. page_number: Page number if available (look for [PAGE X] markers)
            5. regulation_text: The exact text of the statement
            6. statement_type: One of [Obligation, Prohibition, Recommendation, Definition, Exception]
            
            Return ONLY a valid JSON array. Each statement must be self-contained and preserve original text."""
        ).with_model("openai", "gpt-5")
        
        # Truncate if too long (GPT-5 has large context but let's be safe)
        truncated_text = raw_text[:50000] if len(raw_text) > 50000 else raw_text
        
        user_message = UserMessage(
            text=f"""Extract all regulatory statements from this document. Return ONLY a JSON array.
            
Document text:
{truncated_text}

Return format:
[
  {{
    "hierarchy_path": "Chapter 3 > Section 10",
    "section_ref": "10.1",
    "section_title": "Risk Management",
    "page_number": 45,
    "regulation_text": "The exact text...",
    "statement_type": "Obligation"
  }}
]"""
        )
        
        response = await chat.send_message(user_message)
        
        # Parse JSON response
        try:
            # Extract JSON from response
            response_text = response.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:]
            if response_text.startswith('```'):
                response_text = response_text[3:]
            if response_text.endswith('```'):
                response_text = response_text[:-3]
            
            statements = json.loads(response_text.strip())
            return statements
        except json.JSONDecodeError as je:
            logging.error(f"JSON parse error: {je}")
            # Fallback: create basic statements from paragraphs
            return create_fallback_statements(raw_text)
    except Exception as e:
        logging.error(f"AI parsing error: {e}")
        return create_fallback_statements(raw_text)

def create_fallback_statements(text: str) -> List[Dict[str, Any]]:
    """Fallback parser if AI fails"""
    statements = []
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip() and len(p.strip()) > 20]
    
    for idx, para in enumerate(paragraphs[:100], 1):  # Limit to 100 statements
        statements.append({
            "hierarchy_path": f"Section {idx}",
            "section_ref": str(idx),
            "section_title": None,
            "page_number": None,
            "regulation_text": para,
            "statement_type": "Definition"
        })
    
    return statements

# API Endpoints
@api_router.get("/")
async def root():
    return {"message": "Regulation Statement Extractor API"}

@api_router.post("/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """Upload and parse a regulatory document"""
    try:
        # Validate file type
        file_extension = Path(file.filename).suffix.lower()
        allowed_extensions = [".pdf", ".docx", ".html", ".txt"]
        
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"File type {file_extension} not supported")
        
        # Read file content
        file_content = await file.read()
        file_size = len(file_content)
        
        # Create document record
        doc_id = str(uuid.uuid4())
        document = Document(
            id=doc_id,
            filename=file.filename,
            file_type=file_extension,
            file_size=file_size,
            parse_status="processing",
            file_content_base64=base64.b64encode(file_content).decode('utf-8')
        )
        
        doc_dict = document.model_dump()
        doc_dict['upload_date'] = doc_dict['upload_date'].isoformat()
        
        await db.documents.insert_one(doc_dict)
        
        # Extract text based on file type
        raw_text = ""
        page_count = 0
        
        if file_extension == ".pdf":
            raw_text, page_count = await extract_text_from_pdf(file_content)
        elif file_extension == ".docx":
            raw_text, _ = await extract_text_from_docx(file_content)
        elif file_extension == ".html":
            raw_text, _ = await extract_text_from_html(file_content)
        else:  # .txt
            raw_text = file_content.decode('utf-8', errors='ignore')
        
        # Parse with AI
        statements_data = await parse_with_ai(raw_text, doc_id)
        
        # Save statements to database
        statements = []
        for stmt_data in statements_data:
            statement = Statement(
                document_id=doc_id,
                hierarchy_path=stmt_data.get('hierarchy_path', ''),
                section_ref=stmt_data.get('section_ref', ''),
                section_title=stmt_data.get('section_title'),
                page_number=stmt_data.get('page_number'),
                regulation_text=stmt_data.get('regulation_text', ''),
                statement_type=stmt_data.get('statement_type', 'Definition')
            )
            stmt_dict = statement.model_dump()
            stmt_dict['created_at'] = stmt_dict['created_at'].isoformat()
            statements.append(stmt_dict)
        
        if statements:
            await db.statements.insert_many(statements)
        
        # Update document status
        await db.documents.update_one(
            {"id": doc_id},
            {"$set": {
                "parse_status": "completed",
                "total_statements": len(statements)
            }}
        )
        
        return DocumentUploadResponse(
            document_id=doc_id,
            filename=file.filename,
            message=f"Document uploaded and parsed. Extracted {len(statements)} statements."
        )
        
    except Exception as e:
        logging.error(f"Upload error: {e}")
        if 'doc_id' in locals():
            await db.documents.update_one(
                {"id": doc_id},
                {"$set": {"parse_status": "failed"}}
            )
        raise HTTPException(status_code=500, detail=str(e))

@api_router.get("/documents")
async def get_documents():
    """Get all uploaded documents"""
    documents = await db.documents.find({}, {"_id": 0, "file_content_base64": 0}).to_list(1000)
    for doc in documents:
        if isinstance(doc.get('upload_date'), str):
            doc['upload_date'] = datetime.fromisoformat(doc['upload_date'])
    return documents

@api_router.get("/documents/{document_id}/statements")
async def get_statements(document_id: str):
    """Get all statements for a document"""
    statements = await db.statements.find({"document_id": document_id}, {"_id": 0}).to_list(10000)
    for stmt in statements:
        if isinstance(stmt.get('created_at'), str):
            stmt['created_at'] = datetime.fromisoformat(stmt['created_at'])
    return statements

@api_router.put("/statements/{statement_id}")
async def update_statement(statement_id: str, update: StatementUpdate):
    """Update custom fields in a statement"""
    result = await db.statements.update_one(
        {"id": statement_id},
        {"$set": {"custom_fields": update.custom_fields}}
    )
    if result.modified_count == 0:
        raise HTTPException(status_code=404, detail="Statement not found")
    return {"message": "Statement updated successfully"}

@api_router.post("/columns")
async def add_column(column: ColumnDefinition):
    """Add a custom column definition"""
    col_dict = column.model_dump()
    await db.columns.insert_one(col_dict)
    return {"message": "Column added successfully", "column_id": column.id}

@api_router.get("/columns/{document_id}")
async def get_columns(document_id: str):
    """Get all custom columns for a document"""
    columns = await db.columns.find({"document_id": document_id}, {"_id": 0}).to_list(1000)
    return columns

@api_router.delete("/columns/{column_id}")
async def delete_column(column_id: str):
    """Delete a custom column"""
    result = await db.columns.delete_one({"id": column_id})
    if result.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Column not found")
    return {"message": "Column deleted successfully"}

@api_router.post("/templates")
async def save_template(template: Template):
    """Save a column template"""
    template_dict = template.model_dump()
    template_dict['created_at'] = template_dict['created_at'].isoformat()
    await db.templates.insert_one(template_dict)
    return {"message": "Template saved successfully", "template_id": template.id}

@api_router.get("/templates")
async def get_templates():
    """Get all saved templates"""
    templates = await db.templates.find({}, {"_id": 0}).to_list(1000)
    return templates

@api_router.post("/templates/{template_id}/apply")
async def apply_template(template_id: str, document_id: str):
    """Apply a template to a document"""
    template = await db.templates.find_one({"id": template_id}, {"_id": 0})
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    
    # Add columns from template
    for col_data in template['columns']:
        col_data['document_id'] = document_id
        col_data['id'] = str(uuid.uuid4())
        await db.columns.insert_one(col_data)
    
    return {"message": "Template applied successfully"}

@api_router.get("/export/{document_id}")
async def export_data(document_id: str, format: str = Query("csv", regex="^(csv|xlsx)$")):
    """Export statements to CSV or XLSX"""
    # Get statements
    statements = await db.statements.find({"document_id": document_id}, {"_id": 0}).to_list(10000)
    
    if not statements:
        raise HTTPException(status_code=404, detail="No statements found")
    
    # Get custom columns
    columns = await db.columns.find({"document_id": document_id}, {"_id": 0}).to_list(1000)
    
    # Build headers
    base_headers = ["hierarchy_path", "section_ref", "section_title", "page_number", "regulation_text", "statement_type"]
    custom_headers = [col['name'] for col in columns]
    all_headers = base_headers + custom_headers
    
    if format == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(all_headers)
        
        for stmt in statements:
            row = [
                stmt.get('hierarchy_path', ''),
                stmt.get('section_ref', ''),
                stmt.get('section_title', ''),
                stmt.get('page_number', ''),
                stmt.get('regulation_text', ''),
                stmt.get('statement_type', '')
            ]
            for col in columns:
                row.append(stmt.get('custom_fields', {}).get(col['name'], ''))
            writer.writerow(row)
        
        output.seek(0)
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="text/csv",
            headers={"Content-Disposition": f"attachment; filename=export_{document_id}.csv"}
        )
    
    else:  # xlsx
        wb = Workbook()
        ws = wb.active
        ws.title = "Statements"
        
        # Header style
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        
        # Write headers
        for col_idx, header in enumerate(all_headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = header_fill
            cell.font = header_font
        
        # Write data
        for row_idx, stmt in enumerate(statements, 2):
            ws.cell(row=row_idx, column=1, value=stmt.get('hierarchy_path', ''))
            ws.cell(row=row_idx, column=2, value=stmt.get('section_ref', ''))
            ws.cell(row=row_idx, column=3, value=stmt.get('section_title', ''))
            ws.cell(row=row_idx, column=4, value=stmt.get('page_number', ''))
            ws.cell(row=row_idx, column=5, value=stmt.get('regulation_text', ''))
            ws.cell(row=row_idx, column=6, value=stmt.get('statement_type', ''))
            
            for col_idx, col in enumerate(columns, 7):
                ws.cell(row=row_idx, column=col_idx, value=stmt.get('custom_fields', {}).get(col['name'], ''))
        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return StreamingResponse(
            iter([output.getvalue()]),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=export_{document_id}.xlsx"}
        )

@api_router.get("/documents/{document_id}/preview")
async def get_document_preview(document_id: str):
    """Get document preview data"""
    document = await db.documents.find_one({"id": document_id}, {"_id": 0})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "filename": document['filename'],
        "file_type": document['file_type'],
        "file_content_base64": document.get('file_content_base64', '')
    }

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()